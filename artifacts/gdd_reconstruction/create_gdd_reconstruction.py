from __future__ import annotations

from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "纹身构筑_当前策划案重建与可行性审查_v0.1.docx"
FONT_FILE = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_NAME = "Microsoft YaHei"
BLUE = "2E74B5"
NAVY = "1F4D78"
INK = "202B38"
MUTED = "5B6573"
LIGHT = "F4F6F9"
PALE_BLUE = "E8EEF5"
CAUTION = "7A5A00"
RISK = "9B1C1C"
PAGE_WIDTH_DXA = 9360


def cjk_font(size: int, bold: bool = False):
    return ImageFont.truetype(str(FONT_FILE), size=size, index=0)


def draw_round_box(draw, rect, text, fill, outline, font, text_fill=(25, 35, 48), pad=22):
    draw.rounded_rectangle(rect, radius=20, fill=fill, outline=outline, width=3)
    left, top, right, bottom = rect
    max_width = right - left - pad * 2
    lines = []
    for source_line in text.split("\n"):
        current = ""
        for char in source_line:
            candidate = current + char
            if draw.textlength(candidate, font=font) <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = char
        if current:
            lines.append(current)
    bbox = draw.multiline_textbbox((0, 0), "\n".join(lines), font=font, spacing=8, align="center")
    x = left + (right - left - (bbox[2] - bbox[0])) / 2
    y = top + (bottom - top - (bbox[3] - bbox[1])) / 2
    draw.multiline_text((x, y), "\n".join(lines), font=font, fill=text_fill, spacing=8, align="center")


def arrow(draw, start, end, color=(81, 105, 132), width=5):
    draw.line([start, end], fill=color, width=width)
    import math
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    wing = 0.55
    p1 = (end[0] - length * math.cos(angle - wing), end[1] - length * math.sin(angle - wing))
    p2 = (end[0] - length * math.cos(angle + wing), end[1] - length * math.sin(angle + wing))
    draw.polygon([end, p1, p2], fill=color)


def create_mindmaps():
    canvas = (1800, 1060)
    title_font = cjk_font(44, True)
    node_font = cjk_font(28)
    small_font = cjk_font(23)
    title_color = (32, 43, 56)
    bg = (250, 252, 254)
    line = (83, 112, 145)
    maps = []

    # 1. Core loop
    img = Image.new("RGB", canvas, bg)
    d = ImageDraw.Draw(img)
    d.text((80, 55), "思维导图 01｜当前核心循环与玩家决策", font=title_font, fill=title_color)
    center = (690, 410, 1110, 610)
    draw_round_box(d, center, "刺青构筑驱动的\nPvPvE 肉鸽大逃杀", (218, 234, 249), (46, 116, 181), node_font)
    nodes = [
        ((90, 210, 480, 375), "入局选择\n颜料 + 武器 + 图案", (238, 247, 237)),
        ((80, 650, 510, 835), "探索与定位\n地图锚点、NPC、宝箱、事件", (255, 246, 223)),
        ((1310, 190, 1710, 385), "遭遇战\nPVE 敌人、精英、Boss、未来 PVP", (252, 236, 238)),
        ((1305, 650, 1720, 835), "战利品与抉择\n武器、颜料、配方、三选一", (241, 239, 253)),
        ((680, 800, 1120, 980), "缩圈 / 终局\n生存、淘汰与 Run 结算", (234, 245, 247)),
    ]
    anchors = [((480, 300), (690, 450)), ((510, 735), (690, 570)), ((1310, 300), (1110, 450)), ((1305, 740), (1110, 575)), ((900, 800), (900, 610))]
    for a, b in anchors:
        arrow(d, a, b, line)
    for rect, text, fill in nodes:
        draw_round_box(d, rect, text, fill, (120, 135, 150), node_font)
    d.text((100, 930), "注：当前工程已实现 PVE、掉落、NPC、事件与缩圈数据链；真实 PvP 联机仍为待建设方向。", font=small_font, fill=(75, 89, 104))
    p = ROOT / "思维导图_01_核心循环.png"
    img.save(p)
    maps.append((p, "核心循环：从入局构筑、空间决策、战斗到结算的闭环。"))

    # 2. Build system
    img = Image.new("RGB", canvas, bg)
    d = ImageDraw.Draw(img)
    d.text((80, 55), "思维导图 02｜刺青构筑的可读性结构", font=title_font, fill=title_color)
    center = (690, 425, 1110, 640)
    draw_round_box(d, center, "刺青效果\n触发事件 × 元素 × 形状", (218, 234, 249), (46, 116, 181), node_font)
    parts = [
        ((80, 205, 500, 400), "部位（6）\n头：暴击｜躯干：受击\n手臂：技能 / 命中\n腿：闪避 / 移动", (238, 247, 237)),
        ((80, 680, 500, 860), "颜色（7）\n火、雷、自然、冰霜\n变异、神圣、纯净", (255, 246, 223)),
        ((1300, 205, 1720, 400), "图案（8）\n单体、范围、叠层、多段\n链跳、暴击、区域、召唤", (241, 239, 253)),
        ((1300, 680, 1720, 860), "武器（5）\n刀、锤、手枪、弓、能量拳\n提供攻击距离、节奏与瞄准模型", (252, 236, 238)),
    ]
    for rect, text, fill in parts:
        draw_round_box(d, rect, text, fill, (120, 135, 150), node_font)
    for a, b in [((500, 305), (690, 470)), ((500, 770), (690, 600)), ((1300, 305), (1110, 470)), ((1300, 770), (1110, 600))]:
        arrow(d, a, b, line)
    d.text((110, 930), "基础组合：6 个部位 × 7 种颜色 × 8 类图案 = 336 个基础搭配；关键不在数量，而在战斗中能否一眼辨识。", font=small_font, fill=(75, 89, 104))
    p = ROOT / "思维导图_02_刺青构筑.png"
    img.save(p)
    maps.append((p, "刺青构筑：部位负责触发，颜色定义元素，图案定义攻击形态，武器提供基础战斗节奏。"))

    # 3. Status and decision map
    img = Image.new("RGB", canvas, bg)
    d = ImageDraw.Draw(img)
    d.text((80, 55), "思维导图 03｜现状、目标与重构决策边界", font=title_font, fill=title_color)
    center = (650, 425, 1150, 635)
    draw_round_box(d, center, "当前项目\n从 2.5D 原型走向 PvPvE 产品", (218, 234, 249), (46, 116, 181), node_font)
    items = [
        ((75, 190, 520, 375), "已实现 / 可验证\n本地战斗、刺青构筑、敌人、掉落\nNPC、宝箱、三选一、PCG 视觉地图", (238, 247, 237)),
        ((75, 700, 520, 890), "已识别问题\n地图物件不阻挡移动/视线\n视觉摆件多、空间玩法弱、风格混杂", (255, 246, 223)),
        ((1280, 190, 1730, 375), "新版产品目标\n构筑战斗为核心\nPvPvE 大逃杀肉鸽\n提升美术表现与沉浸感", (241, 239, 253)),
        ((1280, 700, 1730, 890), "必须先决策\n镜头（2.5D / 第三人称 / 第一人称）\n局规模、网络权威、场景空间规则", (252, 236, 238)),
    ]
    for rect, text, fill in items:
        draw_round_box(d, rect, text, fill, (120, 135, 150), node_font)
    for a, b in [((520, 285), (650, 470)), ((520, 795), (650, 600)), ((1280, 285), (1150, 470)), ((1280, 795), (1150, 600))]:
        arrow(d, a, b, line)
    d.text((110, 950), "审查原则：先验证“空间 + 构筑 + 对抗”的可玩性，再决定是否承担第一人称和联机重构成本。", font=small_font, fill=(75, 89, 104))
    p = ROOT / "思维导图_03_现状与决策.png"
    img.save(p)
    maps.append((p, "项目现状与重构边界：已实现的本地原型、已暴露的空间问题，以及新版 PvPvE 目标所需决策。"))
    return maps


def set_run_font(run, size=11, color=INK, bold=None, italic=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_cell_text(cell, text, bold=False, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(str(text))
    set_run_font(r, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], PALE_BLUE)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=NAVY, size=9.5)
    mark_header_row(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, align=None, after=8, before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6 if level == 2 else 4)
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 13 if level == 2 else 12, color=BLUE if level < 3 else NAVY, bold=True)
    return p


def add_callout(doc, title, text, color=CAUTION):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    # The callout is a one-row labelled note. Mark its only row as a header so
    # assistive technology announces the note label before its explanatory text.
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(title + "  ")
    set_run_font(r, size=10.5, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, path, caption, alt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.35))
    shape._inline.docPr.set("descr", alt)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.line_spacing = 1.1
    r = cap.add_run(caption)
    set_run_font(r, size=9.5, color=MUTED, italic=True)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(11)
    for style_name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, NAVY)):
        s = styles[style_name]
        s.font.name = FONT_NAME
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("纹身构筑：当前策划案重建与可行性审查｜内部工作稿")
    set_run_font(r, size=8.5, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("第 ")
    set_run_font(r, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)
    r = fp.add_run(" 页")
    set_run_font(r, size=8.5, color=MUTED)
    doc.core_properties.title = "纹身构筑：当前策划案重建与可行性审查"
    doc.core_properties.subject = "当前游戏内容重建、可行性审查与重构决策基线"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "基于当前工程实现、OpenSpec 归档与会话确认内容整理。"


def build_document():
    maps = create_mindmaps()
    doc = Document()
    configure_document(doc)

    # Cover
    add_para(doc, "策划重建与审查", size=12, color=BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18, before=74)
    add_para(doc, "纹身构筑：当前策划案\n重建与可行性审查", size=28, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_para(doc, "v0.1｜以现有实现为证据的可读化基线文档", size=14, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    add_table(doc, ["文档属性", "内容"], [
        ("版本", "v0.1（重建稿，供设计复盘与方向决策使用）"),
        ("整理日期", str(date.today())),
        ("证据范围", "Assets/Game 当前运行代码与资源、OpenSpec 规格/归档、当前会话确认。"),
        ("状态标识", "已实现 / 已规划 / 待决策三类内容严格区分。"),
    ], [2100, 7260])
    add_callout(doc, "阅读结论", "当前项目的价值核心是“身体部位触发的刺青构筑”，不是单纯的俯视角打怪。现阶段最大缺口是空间玩法与产品边界，而不是简单替换一批美术贴图。", NAVY)
    doc.add_page_break()

    add_heading(doc, "内容导览", 1)
    add_para(doc, "本稿按“先还原、再评估、后决策”的顺序组织：先说明已存在的游戏循环与内容资产，再列出其可行性、体验风险与当前方向的未决事项。")
    add_table(doc, ["章节", "阅读目的"], [
        ("1. 文档边界", "明确哪些是事实、哪些是规划、哪些是新版目标。"),
        ("2. 产品概念与核心循环", "还原玩家在一局中的目标、操作与决策。"),
        ("3. 构筑、战斗与内容", "说明刺青、武器、敌人、地图和交互的当前设计。"),
        ("4. 美术与空间问题", "定位“场景空旷、交互难设计”的结构性原因。"),
        ("5. 可行性审查", "评估当前设计的优势、风险与实施门槛。"),
        ("6. 决策清单", "列出继续重构前必须锁定的产品问题。"),
    ], [2400, 6960])

    add_heading(doc, "1. 文档边界与事实分层", 1)
    add_para(doc, "本文件不是新的最终 GDD，也不替代后续的网络、镜头或美术方向决策。它的职责是把散落在规格、代码和资源中的内容，重组成可审查的产品基线。")
    add_table(doc, ["分层", "本稿中的含义", "代表内容"], [
        ("已实现", "当前运行时代码或资源目录中可确认的能力。", "刺青触发链、5 把武器、敌人、宝箱、NPC、三选一、2.5D 相机、PCG 视觉地图。"),
        ("已规划", "OpenSpec 已记录但可能尚未完全落地的设计。", "15–25 分钟局、50 actor 伪联机、分级 AI、三张主题地图、固定地形路线。"),
        ("新版目标", "本轮会话确认、需要通过重构承接的方向。", "以构筑战斗为核心的 PvPvE 大逃杀肉鸽；提升沉浸感与美术表现。"),
        ("待决策", "会显著改变资源、网络和关卡成本的选择。", "2.5D/第三人称/第一人称镜头、首个 PvP 局规模、联网权威模型。"),
    ], [1500, 3350, 4510])
    add_callout(doc, "重要限制", "当前工程未发现 Netcode、Transport、Photon、Mirror 或 FishNet 依赖，也没有多人运行时代码。因此“PvP 大逃杀”目前是产品目标，不是已具备的实现能力。", RISK)

    add_heading(doc, "2. 产品概念与核心循环", 1)
    add_heading(doc, "2.1 一句话概念", 2)
    add_para(doc, "玩家作为末日实验体，在一局高压生存对抗中，把刺青刻在不同身体部位，以元素颜色和图案形状组合出会影响攻击、技能、闪避、移动与承伤的构筑；在敌对生物、环境、资源点以及未来的其他玩家之间做取舍，争取活到终局。")
    add_heading(doc, "2.2 玩家体验支柱", 2)
    add_table(doc, ["体验支柱", "玩家感受", "现有证据"], [
        ("可见构筑", "每一次选择都应让攻击或生存方式发生肉眼可辨的变化。", "部位、颜色、图案三层构筑；效果链已接入攻击、技能、闪避、移动等事件。"),
        ("高压选择", "拿资源、打精英、靠近事件与避开风险之间持续权衡。", "宝箱、死亡宝箱、商人、纹身师、地图事件和三选一的交互链。"),
        ("可学习的战场", "玩家通过路线、地形与敌人信息改进下一局的决策。", "三主题地图与缩圈数据存在，但空间规则目前不足。"),
        ("公平的对抗可读性", "PVP/PVE 中能快速理解威胁、意图与构筑差异。", "这是新版 PvPvE 目标的硬前提，目前尚未完成网络与视觉规范。"),
    ], [1800, 3700, 3860])
    add_figure(doc, maps[0][0], "图 1. 当前核心循环。", maps[0][1])

    add_heading(doc, "2.3 单局流程（重建版）", 2)
    add_para(doc, "入局时，玩家从可用的颜料、武器和图案中完成初始选择。进入战斗地图后，移动、瞄准、攻击、技能与闪避触发各自对应的刺青槽位。击杀、探索和互动带来武器、配方、治疗、金币或一次性效果；玩家再用这些资源调整构筑并应对更强敌人与缩圈，最终以生存/淘汰/结算结束本局。")
    add_callout(doc, "当前最小闭环", "主菜单 → 角色选择 → 初始三选 → 战斗 HUD → 刺青/暂停/结算。NPC、商店和三选一已在系统与规格中存在，但部分 UI 路径曾被标记为后续接入。", CAUTION)

    add_heading(doc, "3. 刺青构筑系统", 1)
    add_heading(doc, "3.1 结构与玩法角色", 2)
    add_para(doc, "刺青不是外观收集物，而是战斗行为的条件化修饰器。部位决定何时触发，颜色决定元素语义，图案决定攻击形态；同一击中事件会因右臂刺青、元素与图案的不同而变成直伤、范围、多段、链跳、叠层或区域效果。基础组合数为 6 × 7 × 8 = 336。")
    add_table(doc, ["身体部位", "触发事件", "主要维度"], [
        ("头部", "暴击", "暴击与爆发"),
        ("躯干", "受击", "生命、防御与反制"),
        ("左臂", "施放技能", "技能强化"),
        ("右臂", "攻击命中", "武器攻击强化"),
        ("左腿", "闪避", "无敌帧、延迟触发与机动收益"),
        ("右腿", "移动", "移速与移动型效果"),
    ], [2100, 3300, 3960])
    add_table(doc, ["颜色（元素）", "图案（形状）"], [
        ("红：火｜黄：雷｜绿：自然｜蓝：冰霜", "线：单体｜环：范围爆发｜螺旋：叠层｜折线：多段"),
        ("紫：变异｜金：神圣｜白：纯净", "闪电：链跳｜星：概率爆发｜流：区域｜兽：召唤"),
    ], [4680, 4680])
    add_figure(doc, maps[1][0], "图 2. 刺青构筑的输入与输出关系。", maps[1][1])
    add_callout(doc, "设计审查重点", "336 不是卖点本身。卖点应是：玩家能否在 1 秒内看懂某个构筑做了什么、在 1 局内形成明确流派、并在 PvP 环境中识别与反制它。", CAUTION)

    add_heading(doc, "4. 战斗、武器与敌人", 1)
    add_heading(doc, "4.1 战斗输入与节奏", 2)
    add_para(doc, "当前输入服务统一管理移动、攻击、蓄力、技能、交互与暂停；攻击既支持鼠标指向的目标选择，也支持不同武器的半角容错。近战强调扇形与贴身风险，远程强调精准、弹药与蓄力。闪避是独立触发事件，因此可以成为构筑的重要节奏点。")
    add_table(doc, ["武器", "定位", "当前基础特性"], [
        ("短刀", "快速近战", "短距离、高容错、快速斩击。"),
        ("重锤", "重击控制", "高伤害、较慢节奏、蓄力爆发。"),
        ("手枪", "稳定远程", "中距离、弹药限制、精准指向。"),
        ("蓄力弓", "高风险狙击", "长距离、需要蓄力、强倍率。"),
        ("能量拳", "近战特种", "较大近距范围与特殊攻击节奏。"),
    ], [1800, 2200, 5360])
    add_heading(doc, "4.2 敌人与遭遇", 2)
    add_para(doc, "敌人体系已有轻型、精英与 Boss 层级，并定义近战、投射、冲锋、跳跃、光束、扇扫、区域脉冲、召唤、护盾、再生和阶段转换等能力。精英承担高价值武器掉落，Boss 承担终局/主题高压。现有主题命名对应 AI 废墟、异星蜂巢与病毒沼泽。")
    add_callout(doc, "PvPvE 风险", "当加入真实玩家后，敌人与 Boss 不能只作为血量更高的 PVE 单位；它们必须创造可以被双方利用、抢夺或规避的空间局面，否则 PvP 与 PvE 会成为两套互不相干的玩法。", RISK)

    add_heading(doc, "5. 地图、空间与交互内容", 1)
    add_heading(doc, "5.1 当前地图内容", 2)
    add_para(doc, "运行时地图当前为 400m 规模，包含 AI_RUINS、ALIEN_HIVE、VIRUS_SWAMP 三个主题。PCG 管线可生成地表切片、装饰物与事件锚点，并为出生、纹身师、商人、Boss 等位置提供布局数据。")
    add_table(doc, ["主题", "氛围定位", "当前空间能力"], [
        ("AI 废墟", "机械遗迹、冷却液、服务金属与失控设施。", "视觉 tile、装饰物、地标与锚点。"),
        ("异星蜂巢", "甲壳、树脂、膜与酸液的生物入侵。", "视觉 tile、装饰物、地标与锚点。"),
        ("病毒沼泽", "腐化、草地、泥地与水域的生化扩散。", "视觉 tile、装饰物、地标与锚点。"),
    ], [1800, 3800, 3760])
    add_heading(doc, "5.2 已定位的空间问题", 2)
    add_para(doc, "当前 PCG 视觉物件在数据层明确被标记为不阻挡移动、也不阻挡视线；地图格默认可走。因此地图可以“摆得更满”，却无法稳定地产生掩体、绕行、卡位、伏击、视线切断或明确的遭遇战边界。这个结构性问题比单纯增加场景贴图更直接地解释了“场景空旷、交互点难设计”。")
    add_callout(doc, "重构原则", "场景资源必须同时回答“它改变了哪一条移动、视线、风险或收益规则”。若资源只是背景装饰，它只能改善截图，不能改善关卡与 PvPvE 的可玩性。", NAVY)
    add_heading(doc, "5.3 互动与经济节点", 2)
    add_table(doc, ["节点", "当前交互", "给核心循环的价值"], [
        ("死亡宝箱", "在死亡 actor 附近拾取金币、颜料、配方或装备。", "将击杀转化为接近风险与战利品选择。"),
        ("武器掉落", "接近后交互拾取；同类可升级，满级转金币。", "让精英与路线选择具有即时构筑价值。"),
        ("宝箱", "可给奖励、治疗或武器。", "提供探索回报与节奏缓冲。"),
        ("纹身师 / 商人", "进入纹身或商店 UI。", "把资源转化为构筑强化。"),
        ("地图事件", "靠近锚点触发三选一。", "提供风险/收益的离散决策。"),
    ], [1800, 3600, 3960])

    add_heading(doc, "6. 视觉现状与美术资源审查", 1)
    add_para(doc, "资源清点显示：当前 Assets/Game 下有约 811 张 PNG、44 个动画文件、2 个 Animator Controller、5 个材质文件，未见可直接复用的 FBX/OBJ 3D 模型。主角已有大量四方向帧动画，Boss 与主题 PCG 资源也以 Sprite 为核心。")
    add_table(doc, ["观察", "影响"], [
        ("整体为 2D / 2.5D 资产管线", "转向第一人称不是替换相机，而是重新建立模型、PBR 材质、碰撞、LOD、第一人称手臂/武器、360° 敌人与关卡生产链。"),
        ("环境大件偏 3D 渲染质感，角色/HUD 偏 2D", "风格层级不统一会削弱世界可信度；需要先建立统一的场景尺度、光照、材质与焦点规则。"),
        ("刺青是身体构筑的视觉载体", "第一人称能突出手臂纹身，但会隐藏躯干、腿部与对手身上的构筑信息；需要额外 UI/特效语言补偿。"),
        ("场景物未产生空间规则", "即使增加 2D 资源，若不新增遮挡/阻挡/地形效果与 POI 模板，空旷感仍会持续。"),
    ], [3200, 6160])
    add_figure(doc, maps[2][0], "图 3. 现有原型、产品目标与重构边界。", maps[2][1])

    add_heading(doc, "7. 可行性与游戏性初审", 1)
    add_heading(doc, "7.1 值得保留的核心", 2)
    add_para(doc, "刺青构筑与身体部位触发具有清晰的差异化：它把传统肉鸽的“数值词条”转化为可表演、可叙事、可被对手观察的身体魔改。武器、技能、闪避和移动都能接入构筑，因此具备形成多流派的基础。宝箱、NPC、事件与精英掉落也已经为一局内的风险/收益循环提供了骨架。")
    add_heading(doc, "7.2 当前最危险的断层", 2)
    add_table(doc, ["断层", "为什么会伤害游戏性", "优先处理方向"], [
        ("空间规则缺失", "战斗发生在没有掩体、路径和视线差异的平面上，策略无法落地。", "先做可阻挡/可遮挡/可造成地形效果的 POI 与遭遇模板。"),
        ("构筑可读性不足", "336 种组合会退化为隐形乘数，玩家无法形成记忆与反制。", "规定每种元素/形状的颜色、动作、命中、范围和音效语义。"),
        ("PvP 目标与本地架构脱节", "没有网络、服务器权威、同步、反作弊和匹配，无法直接验证大逃杀公平性。", "先锁定小规模可玩验证，再建立联网技术路线。"),
        ("镜头选择未决", "镜头决定资源、关卡尺度、敌我信息密度、操作与联机公平性。", "在单个垂直切片内比较 2.5D、第三人称、第一人称，而不是先全量重做。"),
    ], [1900, 3880, 3580])
    add_heading(doc, "7.3 对第一人称的阶段性判断", 2)
    add_para(doc, "第一人称能提升近距离压迫感和武器手感，并让手臂刺青成为强视觉焦点；但它不会自动解决场景设计问题，也不会降低总美术成本。对于以多部位构筑和多人对抗为卖点的产品，第一人称还会降低玩家对自身与他人构筑的可见度。现阶段更合理的结论是：先把“空间规则 + 构筑反馈 + 小规模对抗”做成可测试的垂直切片，再用玩家体验数据决定镜头，而不是把第一人称视为美术问题的直接解法。")

    add_heading(doc, "8. 方向决策清单（进入大重构前）", 1)
    add_para(doc, "以下问题会改变关卡、美术、动画、网络与性能的基本成本，必须在新的总 GDD 中显式定案。")
    add_table(doc, ["决策", "推荐先验", "为何必须先定"], [
        ("首个可玩局规模", "4–8 人等效局（真实玩家或高质量 Bot）", "这是验证构筑对抗、PvE 干扰和缩圈节奏的最低成本范围。"),
        ("镜头", "保留 2.5D 作为对照；以垂直切片验证第三/第一人称", "决定资产格式、场景尺度、信息可读性与玩家操作。"),
        ("地图原则", "固定主地形 + 随机交互锚点；空间规则优先", "有利于记忆、对抗公平与可控的美术生产。"),
        ("刺青可读性", "每个构筑必须有明确的视觉/音效/命中语言", "这是 PvP 反制与产品差异化的共同基础。"),
        ("网络路线", "服务端权威的独立设计项；不从本地逻辑直接外推", "影响同步、预测、命中裁定、作弊防护和运维预算。"),
    ], [2100, 3150, 4110])
    add_callout(doc, "下一份文档应回答的问题", "新版 GDD 不应从“改成第一人称”开始，而应从“玩家在 10 分钟内因何做出有意义的构筑、路线与交战选择”开始；镜头和资源管线随后服务于这一答案。", NAVY)

    add_heading(doc, "附录 A：证据来源与复核范围", 1)
    add_para(doc, "本稿基于下列现有文件重建，未对项目代码、规格或资源做任何修改：")
    add_table(doc, ["类别", "代表来源"], [
        ("核心原始共识", "openspec/changes/archive/2026-06-25-05-gdd-v2-full-design-docs/brainstorm.md 与 proposal.md"),
        ("当前玩法规格", "openspec/specs/player-attack-system/spec.md、weapon-pickup/spec.md、camera-system/spec.md、tattoo/spec.md"),
        ("运行时实现", "Assets/Game/Scripts/Runtime/Services/（TotemCombatService、TotemInteractionService、TotemMapService 等）"),
        ("当前配置目录", "Assets/Game/Scripts/Runtime/TotemGameplayCatalog.cs 与 Assets/Game/DataTable/Business/"),
        ("地图与资源管线", "Assets/Game/Scripts/Runtime/PCGMap/、openspec/changes/28-pcg-map-runtime-integration/、资源目录清点。"),
        ("本轮会话确认", "核心取向：构筑战斗型；产品方向：PvPvE 大逃杀肉鸽；美术与镜头方向待评估。"),
    ], [2400, 6960])
    add_para(doc, "结束语：本稿的价值不在于为现有方案背书，而在于把“已有的可玩骨架”和“未来必须承担的重构成本”放进同一张可讨论的地图中。", size=10.5, color=MUTED, italic=True, after=0)

    doc.save(OUT_DOCX)
    return maps


if __name__ == "__main__":
    result = build_document()
    print(OUT_DOCX)
    for path, _ in result:
        print(path)
