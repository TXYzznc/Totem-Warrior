from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "纹身构筑_3D_PvPvE重构共识与美术垂直切片_v0.2.docx"
MAP = ROOT / "思维导图_04_3D重构共识.png"
FONT = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_NAME = "Microsoft YaHei"
INK = "202B38"
NAVY = "1F4D78"
BLUE = "2E74B5"
MUTED = "5B6573"
PALE = "E8EEF5"
LIGHT = "F4F6F9"
GOLD = "7A5A00"
RISK = "9B1C1C"
WIDTH = 9360


def set_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), color)
    props.append(node)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    props = table._tbl.tblPr
    for name, value in (("tblW", sum(widths)), ("tblInd", 120)):
        node = props.first_child_found_in(f"w:{name}")
        if node is None:
            node = OxmlElement(f"w:{name}")
            props.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            props = cell._tc.get_or_add_tcPr()
            node = props.first_child_found_in("w:tcW")
            if node is None:
                node = OxmlElement("w:tcW")
                props.append(node)
            node.set(qn("w:w"), str(widths[index]))
            node.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header(row):
    props = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    props.append(node)


def set_cell(cell, text, bold=False, color=INK, size=9.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_font(run, size, color, bold)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table_geometry(table, widths)
    for i, value in enumerate(headers):
        shade(table.rows[0].cells[i], PALE)
        set_cell(table.rows[0].cells[i], value, True, NAVY, 9.3)
    mark_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            set_cell(cells[i], value)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_text(doc, text="", size=11, color=INK, bold=False, italic=False, align=None, before=0, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run, size, color, bold, italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6 if level == 2 else 4)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 13 if level == 2 else 12, BLUE if level < 3 else NAVY, True)
    return p


def callout(doc, label, body, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table_geometry(table, [WIDTH])
    mark_header(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    first = p.add_run(label + "  ")
    set_font(first, 10.5, color, True)
    second = p.add_run(body)
    set_font(second, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def font(size, bold=False):
    return ImageFont.truetype(str(FONT), size=size, index=0)


def wrapped(draw, rect, text, fill, outline, text_font):
    draw.rounded_rectangle(rect, radius=22, fill=fill, outline=outline, width=3)
    left, top, right, bottom = rect
    max_width = right - left - 44
    lines = []
    for raw in text.split("\n"):
        line = ""
        for char in raw:
            candidate = line + char
            if draw.textlength(candidate, font=text_font) <= max_width:
                line = candidate
            else:
                lines.append(line)
                line = char
        if line:
            lines.append(line)
    text = "\n".join(lines)
    bounds = draw.multiline_textbbox((0, 0), text, font=text_font, spacing=8, align="center")
    x = left + (right - left - (bounds[2] - bounds[0])) / 2
    y = top + (bottom - top - (bounds[3] - bounds[1])) / 2
    draw.multiline_text((x, y), text, font=text_font, fill=(32, 43, 56), spacing=8, align="center")


def arrow(draw, a, b):
    draw.line([a, b], fill=(83, 112, 145), width=5)
    import math
    angle = math.atan2(b[1] - a[1], b[0] - a[0])
    length = 18
    p1 = (b[0] - length * math.cos(angle - 0.55), b[1] - length * math.sin(angle - 0.55))
    p2 = (b[0] - length * math.cos(angle + 0.55), b[1] - length * math.sin(angle + 0.55))
    draw.polygon([b, p1, p2], fill=(83, 112, 145))


def create_map():
    image = Image.new("RGB", (1800, 1060), (250, 252, 254))
    draw = ImageDraw.Draw(image)
    draw.text((72, 48), "思维导图 04 - 3D PvPvE 重构共识", font=font(44, True), fill=(32, 43, 56))
    center = (655, 420, 1145, 630)
    wrapped(draw, center, "明亮卡通化的\n纹身构筑 PvPvE 撤离大逃杀", (218, 234, 249), (46, 116, 181), font(30, True))
    nodes = [
        ((74, 185, 500, 365), "局内循环\n20 人；8 分钟且剩 40% 开撤离\n可撤离，也可继续争夺", (238, 247, 237)),
        ((75, 690, 510, 890), "纹身构筑\n六部位 × 七色 × 八图案\n工作台读条；成功才覆盖", (255, 246, 223)),
        ((1295, 185, 1725, 365), "局外留存\n八图案熟练度；样式权限\n有限的效率与效果成长", (241, 239, 253)),
        ((1285, 690, 1735, 890), "美术垂直切片\n一名角色 + 一栋可进入建筑\n工作台 UI + 统一灯光", (252, 236, 238)),
        ((680, 800, 1120, 985), "全局美术语言\n明亮活性颜料；局部机械改造\n现代 PC 键鼠 UI", (234, 245, 247)),
    ]
    anchors = [((500, 275), (655, 455)), ((510, 780), (655, 590)), ((1295, 275), (1145, 455)), ((1285, 780), (1145, 590)), ((900, 800), (900, 630))]
    for a, b in anchors:
        arrow(draw, a, b)
    for rect, text, fill in nodes:
        wrapped(draw, rect, text, fill, (120, 135, 150), font(27))
    draw.text((95, 950), "注：首个切片以画面效果为优先，但所有资产仍保持可模块化、可优化的生产结构。", font=font(22), fill=(75, 89, 104))
    image.save(MAP)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(11)
    for name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, NAVY)):
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("纹身构筑：3D PvPvE 重构共识与美术垂直切片 | 内部设计基线")
    set_font(run, 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("第 ")
    set_font(run, 8.5, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    run = footer.add_run(" 页")
    set_font(run, 8.5, MUTED)
    doc.core_properties.title = "纹身构筑：3D PvPvE 重构共识与美术垂直切片 v0.2"
    doc.core_properties.author = "Codex"


def build():
    create_map()
    doc = Document()
    configure(doc)

    add_text(doc, "产品与美术重构基线", 12, BLUE, True, align=WD_ALIGN_PARAGRAPH.CENTER, before=72, after=16)
    add_text(doc, "纹身构筑：3D PvPvE 重构共识\n与美术垂直切片", 28, NAVY, True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_text(doc, "v0.2 | 基于本轮设计共识的执行基线", 14, MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_table(doc, ["项目", "确认内容"], [
        ("产品定位", "刺青构筑驱动的 PvPvE 撤离式大逃杀 Roguelite。"),
        ("首版范围", "PC；20 人为地图与节奏设计标准；50 人不进入当前范围。"),
        ("镜头与地编", "FPS 为主、可切越肩第三人称；所有地编以越肩第三人称为基准。"),
        ("美术目标", "明亮、卡通化、活性颜料世界；效果优先但资产结构必须可优化。"),
        ("验收", "一段 60 秒可玩美术垂直切片，证明角色、建筑、刺青、UI 与灯光的统一性。"),
    ], [2100, 7260])
    callout(doc, "版本定位", "本文件替代 v0.1 中已经被本轮讨论明确推翻或收敛的镜头、局内循环、刺青交互和美术方向假设。它是后续 OpenSpec 规划与垂直切片生产的设计输入，不是立即改造 Unity 代码的实施指令。")
    doc.add_page_break()

    heading(doc, "1. 已确认的产品骨架")
    add_text(doc, "核心体验不是预先带入一套完整答案，而是玩家带着武器与颜料方向进入局内，在随机资源、PVE、玩家威胁与撤离选择中形成当局构筑。局外层保留的是图案掌握、身份展示和有限的熟练优势，而非完整局内刺青。")
    add_table(doc, ["层级", "玩家获得", "失败/结算规则"], [
        ("局内刺青构筑", "六部位上的颜色、图案与触发联动。", "死亡或撤离后均清空；它只属于本局。"),
        ("可掉落物资", "颜料、一次性样式及其他局内物资。", "死亡时由死亡箱承接可掉落部分。"),
        ("全局资源", "成功撤离后才进入局外仓库的材料与收藏。", "未成功撤离不结算入库。"),
        ("图案熟练度", "八种图案的长期掌握、权限与有限成长。", "成功纹身即增加；死亡后保留。"),
    ], [1850, 3650, 3860])
    callout(doc, "长期留存结论", "玩家的长期身份应是“我擅长把随机局面导向某种招牌构筑”，并通过刺青外观、构筑档案和图案熟练度展示出来；不能把最强构筑作为可一键复制的局外预设。")

    heading(doc, "2. 对局、撤离与风险经济")
    add_table(doc, ["规则", "确认值", "设计含义"], [
        ("标准人数", "20 名玩家", "首版地图、资源密度与地编只为 20 人服务。"),
        ("撤离开启", "开局至少 8 分钟，且场内未死亡、未撤离人数不高于初始人数的 40%（20 人局即 8 人）。", "避免过早跳过中期构筑，同时让残局有撤离而非强制决斗的选择。"),
        ("时长", "平均整局约 20 分钟；个人平均 10-15 分钟；整局最长 30 分钟。", "28 分钟左右用毒圈、终局 Boss 或超级补给强制收束。"),
        ("撤离交互", "公开读条；离开范围停止；正常承伤；死亡立即结束并走标准死亡流程。", "撤离是可被埋伏与反制的空间目标。"),
        ("晚留收益", "通过阶段性补给、精英目标或高价值事件增加机会，而非直接结算倍率。", "把“再赌一次”变成可见、可争夺的战术选择。"),
    ], [1500, 3550, 4310])

    heading(doc, "3. 纹身构筑与工作台完整流程")
    add_text(doc, "基础组合沿用既有结构：六个部位决定触发时机，七种颜料决定元素方向，八个图案决定效果形态。总组合为 6 × 7 × 8 = 336；设计目标是可读、可反制，而不是让玩家记忆 336 条独立规则。")
    add_table(doc, ["部位", "触发语义", "构筑职责"], [
        ("脑袋", "暴击命中", "精准、暴击与爆发。"),
        ("躯干正面", "受到伤害", "生存、反击与防御。"),
        ("左臂 / 右臂", "释放技能 / 攻击命中", "技能强化 / 武器攻击形态。"),
        ("左腿 / 右腿", "闪避 / 移动", "机动、闪避与路径控制。"),
    ], [1700, 3100, 4560])
    add_text(doc, "工作台状态：正常操作 → 打开纹身工作台 → 选择部位、图案与颜料并预览 → 纹身读条 → 成功或主动取消 → 玩家主动关闭工作台后恢复操作。")
    add_table(doc, ["规则", "确认行为"], [
        ("工作台打开", "解除鼠标锁定；玩家不能移动、瞄准、攻击、施放技能或进行其他角色控制。"),
        ("读条", "受伤不自动中断；只有 X 主动取消或死亡结束本次纹身。工作台在成功后保持打开。"),
        ("取消", "保留旧刺青；不消耗样式次数；损耗一定比例的颜料。"),
        ("成功", "消耗颜料与一次性样式次数；成功才原子覆盖同部位旧刺青；立即刷新效果、3D 映射与熟练度。"),
        ("死亡", "终止读条；先处理失败损耗，再将剩余可掉落物快照写入死亡箱；刺青构筑清空。"),
    ], [2100, 7260])
    callout(doc, "实现时序护栏", "工作台输入应是独立输入上下文，而非分散在移动、攻击与技能脚本中的禁用判断。纹身读条只能产生一个终态，避免死亡与读条完成同帧时重复消耗、错误覆盖或重复掉落。", RISK)

    heading(doc, "4. 图案熟练度：长期成长但不过度替代局内构筑")
    add_text(doc, "熟练度的单位是图案本身，不区分部位与颜色，因此只有 8 条长期精通路线。初始永久可用图案固定为“直线”和“环形”；其余图案可先以有限次数样式物资使用，再通过成功纹身推进熟练度。")
    add_table(doc, ["阶段", "玩家价值", "建议效果边界"], [
        ("未掌握", "可通过局内有限次数样式试用图案。", "每次成功纹身计入对应图案熟练度。"),
        ("一阶段", "获得长期样式权限。", "在工作台中可使用该图案，不再需要消耗样式次数；仍须消耗颜料并承担读条风险。"),
        ("二阶段", "提高招牌流派的局内构筑效率。", "小幅纹身速度或颜料消耗优化；用加法和总上限控制。"),
        ("三阶段", "有限强化与身份展示。", "小幅图案效果或偏 PVE 的特化收益，并解锁外观、档案和挑战展示。"),
    ], [1600, 3100, 4660])
    callout(doc, "平衡原则", "允许熟练度对 PvP 产生有限、可控影响，但单个阶段收益必须小于一次优秀的局内构筑选择；默认的两种基础图案必须完整可玩，后续图案提供策略广度而非补齐基础能力。", GOLD)

    heading(doc, "5. 全局美术语言")
    add_text(doc, "全局美术不从某张地图出发，而从“明亮卡通化的活性颜料世界”出发。地图只改变建筑与生活场景，不改变角色、刺青、UI、武器与 VFX 的基本语法。危险由高饱和颜料失控、空间收束和资源争夺表现，不依赖低照度或血腥。")
    add_table(doc, ["系统", "全局视觉合同"], [
        ("角色", "局部功能性机械改造的人类；机械接口只框住六处刺青画布，不遮挡刺青。第一版仅一套中性基础身体、统一骨骼和六块固定 UV 区。"),
        ("服装", "轻机能短夹克、背带、护腕、护膝和分段裤；不用厚甲遮住躯干、四肢与头部刺青区。"),
        ("刺青", "哑光合成真皮上的高可读图案；触发时由局部发光、能量流、粒子和音效补足远距离辨识。"),
        ("UI", "现代、干净、面向 PC 键鼠；深色或中性色底，七色颜料只承担元素语义与高价值反馈。"),
        ("VFX / 材质", "圆角硬表面、克制的高光和局部活性颜料发光；避免全屏霓虹使刺青失去焦点。"),
    ], [1900, 7460])

    heading(doc, "6. 第一张地图：生活化未来街区")
    add_text(doc, "首图定位为遭受活性颜料泄漏的未来生活街区。它保留公寓、便利店、纹身店、诊所、咖啡馆、快递柜、地铁口和屋顶花园等日常功能，因此天然支持可进入建筑和室内探索；泄漏后的颜料蔓延、失控设施与 NPC 巢穴再提供风险。")
    add_table(doc, ["地编原则", "对白盒与正式资产的要求"], [
        ("以越肩视角为基准", "掩体、门洞、窗台、转角和室内层高以第三人称可用性验证；第一人称是更沉浸的可切换视角。"),
        ("一个建筑战斗单元", "先做外部接近、至少两条进入路线、室内视线切断、转移/上层路线、另一出口与资源或 PVE 节点。"),
        ("模块化生产", "先建立墙、门、窗、楼板、楼梯、屋顶、隔断与道具套件，再组装生活场景；不从零制作整栋独特建筑。"),
        ("生活逻辑 + 污染逻辑", "每个资产回答“泄漏前做什么、泄漏后如何被颜料改变”；避免无功能的随机赛博装饰。"),
    ], [2200, 7160])

    heading(doc, "7. 60 秒美术垂直切片")
    add_text(doc, "目的不是验证全图规模或最终性能，而是证明明亮卡通化风格能同时支撑沉浸、空间、构筑展示和现代 UI。当前阶段画面效果优先，但所有资产必须保留可模块化、可做 LOD、可烘焙灯光和可替换材质的生产结构。")
    add_table(doc, ["时间", "玩家经历", "必须证明"], [
        ("0-15 秒", "晴天街区移动，观察建筑外观、污染提示和远处目标。", "明亮环境仍有明确危险与探索焦点。"),
        ("15-30 秒", "进入一栋具有上下层或侧路的建筑，遭遇颜料污染与 NPC。", "可进入室内不是装饰；空间能产生视线、掩体与路径选择。"),
        ("30-45 秒", "获得颜料/样式，打开纹身工作台，在 3D 预览中选择并开始纹身。", "UI、角色预览、资源语义与世界材质属于同一风格。"),
        ("45-60 秒", "纹身完成，切换第三人称查看角色与刺青效果，再回到战斗。", "六部位刺青在角色模型、VFX 和战斗反馈中自然成立。"),
    ], [1400, 4050, 3910])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    picture = p.add_run().add_picture(str(MAP), width=Inches(6.35))
    picture._inline.docPr.set("descr", "3D PvPvE 重构共识思维导图")
    add_text(doc, "图 1. 产品、纹身、局外成长与美术垂直切片之间的确认关系。", 9.5, MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    heading(doc, "8. 规模评估与下一阶段")
    add_table(doc, ["领域", "规模判断", "当前处理策略"], [
        ("产品与数值", "大改：局内循环、撤离、掉落、熟练度与经济都需要新规格。", "先把本文件确认的规则拆为可验证的系统规格，再进入数值与掉落表。"),
        ("客户端与联机", "大改：现项目为 2.5D 本地原型，未来 PvPvE 需要独立联机权威设计。", "不从现有本地服务直接外推；先制作离线美术/交互切片。"),
        ("美术与地编", "大改：需建立 3D 角色、模块化建筑、室内空间、UI 和统一材质生产链。", "先完成一名角色、一栋建筑单元、一张工作台 UI 和统一灯光。"),
        ("性能", "暂不以性能 KPI 压缩视觉目标，但不能留下不可优化资产债。", "资产从第一天起保持模块化、可 LOD、可复用材质和可烘焙灯光。"),
    ], [1800, 3650, 3910])
    callout(doc, "OpenSpec 级别结论", "该方向触及镜头、地图、角色、UI、物资、局内状态、局外成长以及未来联机边界，应作为独立的重构变更推进。推荐顺序：先完成美术垂直切片与体验审稿，再进入 20 人 PvPvE 的工程、网络与内容拆分。", RISK)
    add_text(doc, f"整理日期：{date.today()}。本文件以用户本轮确认的设计决定为准；旧 GDD 与历史实现仅作为可复用机制和迁移证据。", 10, MUTED, italic=True, after=0)
    heading(doc, "9. 美术资源需求管理规范")
    add_text(doc, "为避免后续自行寻找资源、委托制作资源与 AI 生成资源混入同一目录，本项目以“需求表先行、指定路径入库、验收后更新状态”为统一流程。美术资源需求根目录固定为 artifacts/美术资源需求，并按职责分为模型、UI、通用三个子目录。")
    for item in [
        "模型：角色、武器、NPC、建筑、室内外模块、交互物与其他三维资产。",
        "UI：HUD、纹身工作台、图标、背包/撤离/死亡箱、熟练度与档案界面。",
        "通用：刺青贴花与遮罩、材质/Shader、VFX、灯光/后处理、Trim Sheet 等跨模块资源。",
        "总表字段固定为：编号（从 1 开始）、名称、类型、路径、描述、状态。状态同时记录是否已完成、是否满足需求或是否需要返工。",
        "资源可由开发者自行寻找，也可由 AI/协作制作；无论来源，只有验收通过后才放入需求表指定路径，并把台账状态更新为“已完成-符合”。"
    ]:
        add_text(doc, "• " + item, after=2)
    callout(doc, "台账原则", "资源的“来源”不决定状态；是否符合当前需求、是否已经验收才决定状态。任何替换资源必须保留编号并更新描述或验收备注，避免场景、UI 与角色风格在迭代中失控。", GOLD)
    doc.save(OUT)
    print(OUT)
    print(MAP)


if __name__ == "__main__":
    build()
